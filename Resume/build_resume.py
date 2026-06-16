# Generates an ATS-friendly .docx for Shubham Khade.
# Rules baked in: single column, Calibri, standard headings, real bullets,
# right-aligned dates via tab stops, contact in body (not header/footer),
# no tables / text boxes / images / columns.

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1A, 0x1A, 0x1A)
RIGHT_TAB = Inches(7.0)  # within usable width for 0.75" margins on Letter

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
pf = normal.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.06

for s in doc.sections:
    s.top_margin = Inches(0.6)
    s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)


def _set_space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def name_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    _set_space(p, 0, 1)
    return p


def center(text, size=10.5, bold=False, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    _set_space(p, 0, after)
    return p


def heading(text):
    p = doc.add_paragraph()
    _set_space(p, 9, 2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def role_date(role, date):
    """Left: bold role. Right (tab-aligned): date."""
    p = doc.add_paragraph()
    _set_space(p, 4, 0)
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(role)
    r.bold = True
    r.font.size = Pt(11)
    t = p.add_run("\t" + date)
    t.font.size = Pt(10)
    return p


def subline(text, italic=True):
    p = doc.add_paragraph()
    _set_space(p, 0, 1)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(10)
    return p


def meta(text):
    p = doc.add_paragraph()
    _set_space(p, 0, 1)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    _set_space(p, 0, 1)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def para(text, after=1, size=10.5):
    p = doc.add_paragraph()
    _set_space(p, 0, after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def labeled(label, value):
    p = doc.add_paragraph()
    _set_space(p, 0, 1)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.size = Pt(10.5)
    v = p.add_run(value)
    v.font.size = Pt(10.5)
    return p


# ---------------- HEADER ----------------
name_line("SHUBHAM ANIL KHADE")
center("Generative AI Engineer - Audio, Vision & LLM Systems", size=11, bold=True, after=2)
center("Pune, India | +91 99229 31551 | shubhamkhade1555@gmail.com", size=10, after=1)
center("linkedin.com/in/shubham-khade-ml | github.com/shubhamkhade1555-creator | shubhamkhade.dev", size=10, after=0)

# ---------------- SUMMARY ----------------
heading("Professional Summary")
para(
    "Generative AI Engineer specializing in production-scale Audio, Computer Vision and LLM systems. "
    "Build end-to-end pipelines across Automatic Speech Recognition (ASR), speaker diarization, voice cloning, "
    "real-time computer vision, and firmware-to-cloud IoT telemetry. First-author of a peer-reviewed machine "
    "learning research paper. Strong in Python, PyTorch, TensorFlow, Large Language Models (LLMs) and real-time "
    "data analytics, with a focus on shipping reliable systems to production.",
    after=1,
)

# ---------------- SKILLS ----------------
heading("Technical Skills")
labeled("Programming Languages", "Python, JavaScript, C, C++, SQL, HTML, CSS")
labeled("Machine Learning & Deep Learning", "PyTorch, TensorFlow, scikit-learn, NumPy, Pandas, CNNs, Computer Vision")
labeled("Audio & Video AI", "Whisper (large-v3), Pyannote 3.1, F5-TTS, XTTS-v2, ElevenLabs, Lip-sync, PySceneDetect")
labeled("Generative AI & LLMs", "Hugging Face, Replicate, Together AI, Ollama, GPT, Claude, Gemini, Llama, Qwen, DeepSeek, Prompt Engineering")
labeled("IoT & Embedded", "Arduino, ESP32, MQTT, AWS IoT Core, RS-485, Firmware")
labeled("Data & Cloud", "MySQL, Firebase, AWS EC2, AWS IoT Core")
labeled("Data Science & Analytics", "Time-Series Analysis, Anomaly Detection, Exploratory Data Analysis (EDA), Data Visualization")
labeled("Tools & DevOps", "Git, Docker, Linux, Postman, Jupyter, VS Code")

# ---------------- EXPERIENCE ----------------
heading("Experience")

role_date("Founding AI Engineer (Core team, 1 of 5)", "Nov 2025 - Present")
subline("Framex Studio, Pune, India")
bullet("Building an end-to-end movie-dubbing pipeline across Indian languages: ASR, speaker diarization, translation, voice-cloned TTS and lip-sync for production-ready workflows.")
bullet("Implemented speaker segmentation with Whisper large-v3 and Pyannote 3.1; timbre-preserving voice cloning using F5-TTS and XTTS-v2.")
bullet("Engineered scene detection, noise removal and spectral enhancement for clean, long-form, multi-speaker output.")

role_date("Graduate Engineer Apprentice - Machine Learning & IoT (Internship)", "Feb 2025 - Oct 2025")
subline("KatMine Technology Pvt. Ltd., India")
bullet("Built a real-time data pipeline ingesting multi-sensor data over RS-485, transforming hardware streams for cloud ingestion and live monitoring.")
bullet("Maintained uptime across the firmware-backend-cloud stack, achieving sub-200 ms end-to-end telemetry latency.")

# ---------------- PROJECTS ----------------
heading("Projects")

role_date("USYNC Audio - AI-Powered Audio Synchronization", "2026")
meta("Live: usync-audio-nu5r.vercel.app  |  Code: github.com/shubhamkhade1555-creator/usync-audio")
bullet("Built and shipped a browser-based tool for AI-powered audio synchronization, aligning and processing audio tracks entirely client-side.")
bullet("Developed end-to-end with JavaScript and the Web Audio API; open-sourced and deployed on Vercel.")

role_date("Real-Time EV Digital Twin & Battery Management System (BMS)", "2024 - 2025")
meta("Live: temperature-demo.vercel.app  |  Code: github.com/shubhamkhade1555-creator/Temperature_demo")
bullet("Designed a real-time telemetry pipeline ingesting 5+ sensor streams (voltage, current, temperature, RPM, GPS) into a cloud analytics layer.")
bullet("Implemented cell-level anomaly detection, driver-behaviour scoring, fleet-health analytics and predictive maintenance over an EV digital twin.")
bullet("Stack: Python, Pandas, Time-Series Analysis, Anomaly Detection, MQTT, AWS IoT Core.")

role_date("ML Sign-Language System for the Deaf (Published Research)", "2025")
meta("Code: github.com/shubhamkhade1555-creator/ml-sign-language-system-for-the-deaf  |  Paper: GRENZE GIJET Vol 11(2), 2025")
bullet("Developed real-time CNN gesture recognition translating sign language to text from video frames (OpenCV + TensorFlow).")
bullet("Introduced a novel haptic vibration-feedback channel enabling two-way deaf-to-hearing communication.")

para("Additional Projects:", after=1, size=10.5)
bullet("CCTV-Based Early Fire Detection - real-time fire and smoke detection on existing CCTV feeds with automated alerts (Python, OpenCV, CNN).")
bullet("PPE Safety Compliance System - real-time safety-shoe detection on shop floors for automated compliance checks (Python, YOLO, OpenCV).")
bullet("ApplyPilot AI - autonomous job-application agent running Gemma 2B locally for job-description parsing and relevance scoring (Python, LLM, Selenium).")
bullet("USYNC - AI design-system and artifact hub indexing 150+ AI-generated artifacts with live previews (Prompt Engineering, Vanilla JS, Vercel).")

# ---------------- PUBLICATIONS ----------------
heading("Publications")
bullet('Khade, S. (First Author). "Intelligent Machine Learning Communication System for the Deaf." GRENZE Journal of Engineering & Technology (GIJET), Vol 11, Issue 2, 2025 (Paper ID 6348).')

# ---------------- CERTIFICATIONS ----------------
heading("Certifications")
bullet("AWS Academy Graduate - Getting Started with AWS IoT, AWS Training & Certification, 2025.")
bullet("Introduction to Amazon EC2, AWS Training & Certification, 2026.")

# ---------------- ACHIEVEMENTS ----------------
heading("Achievements")
bullet('Presented "IoT-Enabled Two-Wheeler Accident Detection System" at the State-Level IoT Project Competition (2024-25).')

# ---------------- EDUCATION ----------------
heading("Education")
role_date("B.Tech, Computer Science & Engineering (IoT & Cyber Security with Blockchain)", "2021 - 2025")
meta("Annasaheb Dange College of Engineering & Technology, Ashta  |  CGPA: 7.1/10")
role_date("Higher Secondary Certificate (Class XII), Science", "2021")
meta("Dr. Bapuji Salunkhe College, Sangli  |  83%")
role_date("Secondary School Certificate (Class X)", "2019")
meta("Vidyamandir Prashala, Miraj  |  73.6%")

out = r"C:\Users\shubh\Desktop\portfolio\Resume\Shubham_Khade_Resume.docx"
doc.save(out)
print("Saved:", out)
